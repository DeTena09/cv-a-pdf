from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from config import IMAGE_PATH

def set_margins(doc, top=2.54, bottom=2.54, left=3, right=3):
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top / 2.54)
        section.bottom_margin = Inches(bottom / 2.54)
        section.left_margin = Inches(left / 2.54)
        section.right_margin = Inches(right / 2.54)

def agregar_titulo(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)

def agregar_subtitulo(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto.upper())
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)

def agregar_linea(doc, texto, negrita=False, salto=True):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrita
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    if not salto:
        p.paragraph_format.space_after = Pt(0)
    return p

def agregar_linea_habilidad(doc, titulo, contenido):
    p = doc.add_paragraph()
    
    # Primer run: el título en negrita
    run1 = p.add_run(f"{titulo}: ")
    run1.bold = True
    run1.font.size = Pt(10)
    run1.font.name = 'Arial'
    
    # Segundo run: el contenido normal
    run2 = p.add_run(contenido)
    run2.font.size = Pt(10)
    run2.font.name = 'Arial'
    
    # Ajustes de formato del párrafo
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = Pt(3)

def eliminar_bordes(tabla):
    tbl = tabla._tbl
    for cell in tbl.iter():
        if cell.tag == qn('w:tcBorders'):
            tbl.remove(cell)

def crear_cv_desde_json(datos, output_path):
    doc = Document()

    # Márgenes
    set_margins(doc)

    # Agregar logo alineado a la izquierda
    p_logo = doc.add_paragraph()
    run_logo = p_logo.add_run()
    run_logo.add_picture(IMAGE_PATH, width=Inches(1.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # === NOMBRE ===
    agregar_titulo(doc, datos["nombre"])

    # === PROFILE ===
    agregar_subtitulo(doc, "PROFILE:")
    for linea in datos["perfil"]:
        linea = linea.strip()
        if not linea:
            continue
        parrafo = doc.add_paragraph(style="List Bullet")
        parrafo.paragraph_format.left_indent = Inches(0.6)
        parrafo.paragraph_format.first_line_indent = Inches(-0.2)
        parrafo.paragraph_format.line_spacing = 1
        run = parrafo.add_run()

        if " (" in linea and linea.endswith(")"):
            texto_principal = linea[:linea.rfind(" (")]
            duracion = linea[linea.rfind(" (") + 2 : -1]
            run1 = parrafo.add_run(texto_principal + " (")
            run2 = parrafo.add_run(duracion)
            run2.bold = True
            run3 = parrafo.add_run(")")
        else:
            parrafo.add_run(linea)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # === EXPERIENCIA PROFESIONAL ===
    agregar_subtitulo(doc, "EXPERIENCIA PROFESIONAL:")
    for exp in datos["experiencia"]:
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = True
        cells = table.rows[0].cells

        eliminar_bordes(table)  # Quitar bordes

        texto_celdas = [
            exp['empresa'],
            exp['localizacion'] or "",
            exp['rol'],
            f"Desde:\n{exp['desde']}",
            f"Hasta:\n{exp['hasta']}"
        ]
        for i, texto in enumerate(texto_celdas):
            par = cells[i].paragraphs[0]
            run = par.add_run(texto)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            par.paragraph_format.space_after = Pt(0)
            par.paragraph_format.space_before = Pt(0)

        
        doc.add_paragraph().paragraph_format.space_after = Pt(14)

       # Añadir pequeña tabla 2 columnas para "Experiencia funcional"
        tabla_exp = doc.add_table(rows=1, cols=2)
        tabla_exp.alignment = WD_TABLE_ALIGNMENT.LEFT
        tabla_exp.autofit = True
        eliminar_bordes(tabla_exp)

        cell1, cell2 = tabla_exp.rows[0].cells

        # Celda de TÍTULO
        p1 = cell1.paragraphs[0]
        run1 = p1.add_run("Experiencia funcional:")
        run1.font.size = Pt(10)
        run1.font.name = 'Arial'

        # Celda de TEXTO
        p2 = cell2.paragraphs[0]
        run2 = p2.add_run(exp["funcional"])
        run2.font.size = Pt(10)
        run2.font.name = 'Arial'
        p2.paragraph_format.line_spacing = 1.5

        # Añadir pequeña tabla 2 columnas para "Herramientas"
        tabla_herr = doc.add_table(rows=1, cols=2)
        tabla_herr.alignment = WD_TABLE_ALIGNMENT.LEFT
        tabla_herr.autofit = True
        eliminar_bordes(tabla_herr)

        cell1h, cell2h = tabla_herr.rows[0].cells

        # Celda de TÍTULO
        p1h = cell1h.paragraphs[0]
        run3 = p1h.add_run("Herramientas:")

        run3.underline = True
        run3.font.size = Pt(10)
        run3.font.name = 'Arial'
        

        # Celda de TEXTO
        p2h = cell2h.paragraphs[0]
        run4 = p2h.add_run(exp["herramientas"])
        run4.font.size = Pt(10)
        run4.bold = True
        run4.font.name = 'Arial'
        p2h.paragraph_format.line_spacing = 1.5

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # === FORMACIÓN ACADÉMICA ===
    agregar_subtitulo(doc, "FORMACIÓN ACADÉMICA:")
    for form in datos["formacion"]:
        agregar_linea(doc, f"{form['titulo']} — {form['centro']} ({form['fecha']})")

    # === CERTIFICACIONES ===
    if "certificados" in datos and datos["certificados"]:
        agregar_subtitulo(doc, "CERTIFICACIONES:")
        for cert in datos["certificados"]:
            agregar_linea(doc, f"{cert['nombre']} – {cert['institucion']} ({cert['fecha']})")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # === IDIOMAS ===
    agregar_subtitulo(doc, "IDIOMAS:")
    for idioma in datos["idiomas"]:
        agregar_linea(doc, f"{idioma['idioma']}: {idioma['nivel']}")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # === HABILIDADES TÉCNICAS ===
    agregar_subtitulo(doc, "HABILIDADES TÉCNICAS:")
    habilidades = datos["habilidades"]
    agregar_linea_habilidad(doc, "Perfil principal", habilidades['perfil_principal'])
    agregar_linea_habilidad(doc, "Perfil secundario", habilidades['perfil_secundario'])
    agregar_linea_habilidad(doc, "Herramientas", habilidades['herramientas'])


    doc.save(output_path)
    return output_path


